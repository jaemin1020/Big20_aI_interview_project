
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font_style(run, name='맑은 고딕', size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def create_ultra_detailed_report():
    doc = Document()
    
    # --- Title Page ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\nAI-Worker 차세대 지능형 면접 엔진\n상세 기술 설계 및 알고리즘 분석 보고서")
    set_font_style(run, size=24, bold=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nAI Interview Project: AI-Worker Part Deep Dive")
    set_font_style(run, size=14, italic=True)
    
    doc.add_page_break()

    # --- 1. Executive Summary ---
    doc.add_heading('1. 서론 (Executive Summary)', level=1)
    doc.add_paragraph(
        "본 보고서는 AI 면접 서비스의 핵심 지능을 담당하는 'AI-Worker' 파트의 기술적 고도화 성과를 정리합니다. "
        "기존의 단순 질문 나열 방식을 넘어, 지원자의 이력서 데이터를 심층 분석(RAG)하고 기업 고유의 인재상을 알고리즘에 실시간으로 반영하는 "
        "'동적 맞춤형 면접 엔진'의 설계와 구현 상세를 기술합니다."
    )
    doc.add_paragraph(
        "핵심 아키텍처는 Retrieval-Augmented Generation (RAG)을 기반으로 하며, EXAONE-3.5 언어 모델을 통해 "
        "인간 면접관에 준하는 맥락 이해와 날카로운 심층 꼬리질문 생성 능력을 갖추었습니다."
    )

    # --- 2. 이력서 파싱 및 구조화 엔진 ---
    doc.add_heading('2. 이력서 파싱 및 구조화 엔진 분석', level=1)
    doc.add_heading('2.1. PDF 레이아웃 분석 및 텍스트 추출', level=2)
    doc.add_paragraph(
        "비정형 데이터인 PDF 이력서에서 의미 있는 정보를 손실 없이 추출하기 위해 PyMuPDF(fitz) 라이브러리를 활용했습니다. "
        "단순한 텍스트 나열이 아닌, 좌표값 기반의 폰트 크기 및 단락 분석을 수행하여 제목과 본문을 구분합니다. "
        "이미지 형태의 스캔 문서에 대해서는 Tesseract 또는 관련 OCR 모듈을 연동할 수 있는 인터페이스를 갖추고 있습니다."
    )
    doc.add_heading('2.2. LLM 기반 Zero-shot 정보 추출 및 구조화', level=2)
    doc.add_paragraph(
        "추출된 원문은 LLM을 통해 정교한 JSON 구조로 변환됩니다. 별도의 라벨링된 학습 데이터 없이도 "
        "LLM의 명령 수행 능력을 활용하여 인적사항(Header), 학력(Education), 경력(Activities), 프로젝트(Projects), 자격증(Certifications) 섹션을 정확히 추출합니다. "
        "이는 정규식이나 고전적인 NLP 기반 방식보다 월등한 유연성과 정확도를 제공합니다."
    )

    # --- 3. RAG 파이프라인 엔진 ---
    doc.add_heading('3. RAG (Retrieval-Augmented Generation) 시스템 설계', level=1)
    doc.add_heading('3.1. 지능형 텍스트 청킹(Chunking) 전략', level=2)
    doc.add_paragraph(
        "검색의 정밀도를 결정하는 청킹 단계에서는 'Recursive Character Text Splitter' 알고리즘을 사용합니다. "
        "단순 길이 기반 분할이 아닌 문단('\n\n'), 문장('\n'), 공백(' ') 순으로 우선순위를 두어 분할함으로써 "
        "하나의 의미 단위가 검색 조각에서 잘려나가는 것을 방지합니다. 또한 각 청크 간 50자 내외의 Overlap 구간을 설정하여 맥락의 연속성을 확보했습니다."
    )
    doc.add_heading('3.2. 고차원 벡터 임베딩 및 pgvector 검색', level=2)
    doc.add_paragraph(
        "HuggingFace의 고성능 임베딩 모델을 활용하여 텍스트를 1024차원의 밀집 벡터(Dense Vector)로 변환합니다. "
        "저장소로는 PostgreSQL의 벡터 익스텐션인 'pgvector'를 채택하여, 기존의 관계형 데이터와 벡터 데이터를 통합 관리하는 효율성을 달성했습니다. "
        "검색 알고리즘으로는 Cosine Similarity를 사용하며, 대규모 데이터셋에서의 검색 속도를 위해 HNSW(Hierarchical Navigable Small World) 인덱싱 기법을 적용했습니다."
    )

    # --- 4. Vector Database 아키텍처 설계 ---
    doc.add_heading('4. Vector Database 및 고속 검색 엔진 설계', level=1)
    doc.add_heading('4.1. pgvector 기반의 통합 데이터 관리', level=2)
    doc.add_paragraph(
        "본 시스템은 검색 성능과 데이터 일관성을 동시에 확보하기 위해 PostgreSQL의 벡터 익스텐션인 'pgvector'를 채택했습니다. "
        "전통적인 벡터 전용 DB(Pinecone, Milvus 등)를 별도로 구축하는 대신, pgvector를 사용함으로써 "
        "지원자의 인적 사항, 면접 기록 등의 관계형 데이터와 이력서 임베딩과 같은 고차원 벡터 데이터를 단일 트랜잭션 범위 내에서 통합 관리할 수 있는 아키텍처를 실현했습니다."
    )
    doc.add_heading('4.2. HNSW 인덱싱 및 Cosine Distance 알고리즘', level=2)
    doc.add_paragraph(
        "수만 개의 이력서 청크와 질문 은행 데이터 사이에서 밀리초(ms) 단위의 검색 속도를 보장하기 위해 'HNSW(Hierarchical Navigable Small World)' 인덱싱 기법을 적용했습니다. "
        "이는 벡터 공간에 그래프 구조를 구축하여 검색 범위를 기하급수적으로 좁히는 알고리즘으로, 정확도 손실을 최소화하면서도 압도적인 검색 성능을 제공합니다."
    )
    doc.add_paragraph(
        "벡터 간의 유사도 측정 방식으로는 'Cosine Distance(1 - Cosine Similarity)'를 채택했습니다. "
        "이는 문서의 길이나 단어 빈도에 민감한 유클리디안 거리보다, 질문과 문맥 사이의 '의미적 방향성'을 더 정확하게 포착하여 최적의 면접 질문 소재를 찾아내는 데 적합합니다."
    )

    # --- 5. 질문 생성 및 시나리오 엔진 ---
    doc.add_heading('5. 실시간 시나리오 엔진 및 질문 생성 알고리즘', level=1)
    doc.add_heading('4.1. 기업 인재상(Talent Image) 동적 주입 시스템', level=2)
    doc.add_paragraph(
        "이번 고도화의 핵심 성과로서, 지원 회사의 고유 인재상을 실시간 질문 생성에 반영하는 로직을 구축했습니다. "
        "DB의 'companies' 테이블에서 해당 기업의 'ideal' 필드를 조회하여 LLM 가이드에 변수({company_ideal}) 형태로 주입합니다. "
        "이를 통해 AI는 단순한 기술 질문을 넘어 그 회사의 비전과 가치관에 지원자가 얼마나 부합하는지를 묻는 '인재상 맞춤형 상황 질문'을 생성합니다."
    )
    doc.add_heading('4.2. Context-Aware Follow-up (심층 꼬리질문) 로직', level=2)
    doc.add_paragraph(
        "지원자가 답변을 하면, AI-Worker는 이전 질문과 지원자의 답변을 비교 분석합니다. "
        "답변 내에서 핵심 키워드(Key-Phrase)를 추출하여 '지원자님의 ~~라는 경험에 대해 좀 더 구체적으로 설명해주세요'와 같이 "
        "답변 내용을 적극적으로 인용(Citation)함으로써 면접의 몰입도와 검증의 깊이를 극대화합니다."
    )

    # --- 5. 분석 및 평가 엔진 ---
    doc.add_heading('5. 면접 평가 및 분석 엔진 (Evaluator)', level=1)
    doc.add_paragraph(
        "면접 종료 후, `evaluator.py` 모듈은 전제 면접 스크립트를 분석하여 다각도 역량 리포트를 생성합니다. "
        "평가 알고리즘은 사전에 정의된 역량 루브릭(Technical, Communication, Cultural Fit)과 지원자의 발화를 대조합니다. "
        "단순한 텍스트 분석을 넘어 감정 분석(Sentiment Analysis)과 비언어적 데이터(Vision)를 통합하여 "
        "지원자의 자신감, 태도 등을 수치화하고 종합적인 피드백을 산출합니다."
    )

    # --- 6. 시스템 최적화 및 인프라 ---
    doc.add_heading('6. 분산 처리 및 성능 최적화 전략', level=1)
    doc.add_paragraph(
        "추론 부하가 큰 LLM 작업을 원활하게 처리하기 위해 Celery와 Redis를 활용한 비동기 작업 큐 시스템을 적용했습니다. "
        "웹 서버(API)와 추론 엔진(Worker)를 분리하여 브라우저 응답 지연을 방지합니다. "
        "또한 GPU 가용 메모리 관리(Torch caching) 및 Garbage Collection을 통해 장시간 면접 중에도 시스템이 중단되지 않는 안정성을 확보했습니다."
    )

    # --- Tables Section ---
    doc.add_page_break()
    doc.add_heading('7. 기술 사양 및 데이터 프로토콜 요약 (Tables)', level=1)

    # Table 1: Modules
    doc.add_heading('7.1. 주요 모듈 및 기술 스택', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = '모듈명', '주요 기능', '적용 기술', '특징'
    
    mod_data = [
        ['Resume Parser', '이력서 구조화', 'fitz, LLM Extraction', 'JSONB 저장'],
        ['RAG Engine', '문맥 정보 추출', 'pgvector, HuggingFace', '1024차원 벡터 검색'],
        ['Question Gen', '맞춤형 질문 생성', 'EXAONE-3.5, prompt engineering', '인재상 동적 반영'],
        ['Evaluator', '종합 역량 평가', 'Rubric Analysis, Sentiment', 'spider chart 데이터 생성'],
        ['TTS / STT', '음성 처리', 'Whisper, Synthesis', '실시간 스트리밍 인터랙션'],
        ['Vision AI', '비언어 분석', 'Pose/Emotion detection', '태도 점수 산출']
    ]
    for m in mod_data:
        row = table.add_row().cells
        for i, val in enumerate(m): row[i].text = val

    # Table 2: RAG Pipeline
    doc.add_heading('7.2. RAG 파이프라인 상세 프로토콜', level=2)
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '공정 단계', '메커니즘', '기대 효과'
    rag_data = [
        ['1. Chunking', 'Recursive Split (Overlap 50자)', '문맥 단절 방지'],
        ['2. Embedding', 'Dense Vector (1024차원)', '의미적 유사도 파악'],
        ['3. Retrieval', 'pgvector Cosine Distance', '고속 관련 정보 추출'],
        ['4. Post-processing', 'LLM Guide Integration', '할루시네이션(환각) 방지']
    ]
    for r in rag_data:
        row = table2.add_row().cells
        for i, val in enumerate(r): row[i].text = val

    # --- Conclusion ---
    doc.add_heading('8. 결론 및 향후 전망', level=1)
    doc.add_paragraph(
        "AI-Worker는 정교한 RAG 알고리즘과 유연한 시나리오 엔진을 결합하여, 지원자와 기업 모두에게 "
        "가장 의미 있는 면접 데이터를 제공하는 독보적인 완성도를 확보했습니다. "
        "향후에는 실시간 데이터 기반의 '적응형 난이도 조절(Adaptive Difficulty)' 기능을 추가하여 기술적 완성도를 더욱 높일 예정입니다."
    )

    file_path = r'C:\big20\Big20_aI_interview_project\AI_Worker_Super_Detailed_Report.docx'
    doc.save(file_path)
    return file_path

if __name__ == "__main__":
    path = create_ultra_detailed_report()
    print(f"Super Detailed Report created at: {path}")
