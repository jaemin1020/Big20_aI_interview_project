
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_detailed_report():
    doc = Document()

    # 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('AI-Worker 면접 엔진 기술 상세 보고서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. 개요
    doc.add_heading('1. 프로젝트 개요 및 목적', level=1)
    doc.add_paragraph(
        "본 프로젝트의 AI-Worker 파트는 면접 시스템의 '지능형 브레인' 역할을 수행하며, 기술적 숙련도 검증을 넘어 지원자의 가치관이 기업의 인재상에 부합하는지를 동적으로 평가하는 것을 목적으로 합니다. "
        "기존의 경직된 정적 시나리오 체계를 탈피하여, 실시간 데이터 연동과 LLM(Large Language Model) 알고리즘을 결합한 '동적 맞춤형 시나리오 엔진'을 구축하였습니다."
    )

    # 2. 모듈별 기술 분석
    doc.add_heading('2. 핵심 모듈별 알고리즘 및 데이터 파이프라인', level=1)

    doc.add_heading('2.1. 이력서 구조화 엔진 (Resume Parsing & Structuring)', level=2)
    doc.add_paragraph(
        "비정형 PDF 데이터로부터 의미 있는 정보를 추출하기 위해 PyMuPDF 기반의 레이아웃 분석 알고리즘을 적용했습니다. "
        "추출된 raw 텍스트는 LLM의 Zero-shot Information Extraction 기법을 통해 학력, 경력, 프로젝트, 자격증 등의 JSON 구조로 변환됩니다. "
        "특히, 기술 스택 정규화 로직을 통해 상이한 명칭들을 시스템 표준 메타데이터로 일원화하여 매칭 정확도를 극대화했습니다."
    )

    doc.add_heading('2.2. RAG (Retrieval-Augmented Generation) 파이프라인', level=2)
    doc.add_paragraph(
        "지원자의 방대한 이력서 중 가장 적합한 문맥을 찾아내기 위해 고도화된 RAG 프로세스를 구현하였습니다."
    )
    doc.add_paragraph(
        "- Chunking Strategy: 문맥 단절을 최소화하기 위해 'Recursive Character Splitting' 방식을 사용하여 500자 단위로 쪼개고, 50자의 중첩(Overlap) 구간을 두어 정보의 연속성을 보장했습니다.\n"
        "- Vector Search: HuggingFace의 1024차원 고밀도 임베딩 모델과 PostgreSQL의 pgvector를 결합했습니다. 'Cosine Similarity' 알고리즘을 기반으로 질문 의도와 가장 유사한 이력서 구절을 TOP-K 방식으로 실시간 추출합니다."
    )

    doc.add_heading('2.3. 지능형 질문 생성 엔진 (Dynamic Question Generation)', level=2)
    doc.add_paragraph(
        "단순한 질문 생성을 넘어, 'Context-Aware Prompting' 기술을 통해 현재 면접의 흐름, 지원자의 이전 답변, 기업의 인재상을 하나의 프롬프트로 결합합니다.\n"
        "- 어조 제어 알고리즘: TTS 서버와의 호환성을 높이기 위해 특수문자 및 불필요한 태그를 자동 클리닝하며, 상황에 맞는 어미(예: ~주세요, ~인가요?)를 유연하게 적용합니다.\n"
        "- 실시간 꼬리질문: 지원자의 답변에서 핵심 명사와 동사를 실시간 추출하여 '...라고 말씀하셨는데'와 같이 인용하며 파고드는 지능형 Follow-up 로직을 설계했습니다."
    )

    # 3. 기업 인재상 동적 연동 프로세스
    doc.add_heading('3. 기업 인재상(Talent Image) 동적 연동 및 검증 전략', level=1)
    doc.add_paragraph(
        "본 시스템의 가장 큰 차별점은 DB에 저장된 실제 기업의 인재상을 면접 시나리오에 실시간 주입한다는 점입니다."
    )
    doc.add_paragraph(
        "1) 데이터 로드: 인터뷰 생성 시 특정된 회사 ID를 기반으로 'companies' 테이블의 'ideal' 필드를 조회합니다.\n"
        "2) 가이드라인 매핑: 조회된 인재상 텍스트를 {company_ideal} 변수를 통해 AI 가이드에 주입합니다.\n"
        "3) 가치 검증 질문 설계: AI는 주입된 인재상을 해석하여 '기술의 민주화', '책임감', '창의적 융합' 등 기업이 중시하는 키워드에 맞춘 상황 질문을 생성합니다. 이를 통해 지원자가 기술적 역량뿐만 아니라 해당 기업의 문화적 적합성(Cultural Fit)을 갖췄는지 정밀 검증합니다."
    )

    # 4. 기술 사양 요약 표
    doc.add_heading('4. 시스템 기술 사양 요약', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '구분'
    hdr_cells[1].text = '핵심 기술'
    hdr_cells[2].text = '적용 알고리즘'
    hdr_cells[3].text = '비고'
    
    specs = [
        ['데이터 처리', 'RAG', 'Cosine Similarity', '이력서 맞춤형 검색'],
        ['LLM 추론', 'EXAONE-3.5', 'Persona-based Prompting', '기업별 인재상 연동'],
        ['데이터 베이스', 'PostgreSQL', 'pgvector (HNSW Index)', '고속 벡터 검색 지원'],
        ['비동기 인프라', 'Celery / Redis', 'Distributed Task Queue', 'API 응답 지연 방지'],
        ['멀티모달', 'STT / TTS', 'Whisper / Stream Synthesis', '실시간 대화 구현']
    ]
    for spec in specs:
        row_cells = table.add_row().cells
        for i, val in enumerate(spec):
            row_cells[i].text = val

    # 5. 기대 효과 및 결론
    doc.add_heading('5. 기대 효과 및 향후 발전 방향', level=1)
    doc.add_paragraph(
        "이번 고도화를 통해 구축된 AI-Worker 엔진은 기업과 지원자 모두에게 최적화된 면접 경험을 제공합니다. 기업은 코드 수정 없이 DB 업데이트만으로 맞춤형 시나리오를 즉시 운영할 수 있으며, 지원자는 자신의 강점과 기업의 비전이 맞닿은 지점에서 진정성 있는 답변을 끌어낼 수 있습니다."
    )
    doc.add_paragraph(
        "향후에는 생성된 질문의 품질을 정량적으로 측정하는 자동 모니터링 기능과 더불어, 기업별 면접 스타일(Persona)을 더욱 세분화하여 서비스의 전문성을 강화할 예정입니다."
    )

    file_path = r'C:\big20\Big20_aI_interview_project\AI_Worker_Detailed_Technical_Report.docx'
    doc.save(file_path)
    return file_path

if __name__ == "__main__":
    path = create_detailed_report()
    print(f"Detailed Report created at: {path}")
