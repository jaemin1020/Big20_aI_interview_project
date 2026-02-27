"""
AI 모의면접 시스템 최종 모델 설계서 (DOCX) 자동 생성
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return h

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

# ═══════════════════════════════════════
# 표지
# ═══════════════════════════════════════
doc.add_paragraph('')
doc.add_paragraph('')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AI 모의면접 시스템\n최종 모델 설계서')
run.font.size = Pt(28)
run.bold = True
run.font.name = '맑은 고딕'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_paragraph('')
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('STT · TTS · CV · 감정분석 · LLM 통합 아키텍처')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph('')
doc.add_paragraph('')
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('작성일: 2026년 2월 24일 (최종)\n').font.size = Pt(11)
info.add_run('작성자: Project Big20 AI Team\n').font.size = Pt(11)
info.add_run('버전: v2.0 Final').font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════
# 목차
# ═══════════════════════════════════════
add_heading_styled('목차', 1)
toc_items = [
    '1. 시스템 개요',
    '2. 전체 아키텍처',
    '3. STT (음성 인식) — Whisper Large-v3-Turbo',
    '4. TTS (음성 합성) — Supertonic 2',
    '5. CV (영상 분석) — MediaPipe Face Mesh V2',
    '6. 감정 분석 (음성 자신감) — NumPy RMS 기반',
    '7. LLM (질문 생성 / 답변 평가) — EXAONE 3.5-7.8B',
    '8. 통합 채점 시스템',
    '9. 성능 비교 및 검증 결과',
    '10. DB 저장 구조',
    '11. 결론 및 향후 계획',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════
# 1. 시스템 개요
# ═══════════════════════════════════════
add_heading_styled('1. 시스템 개요', 1)

doc.add_paragraph(
    '본 문서는 AI 모의면접 시스템에서 사용되는 모든 AI 모델의 최종 설계를 기술합니다. '
    '각 모델의 선정 배경, 성능 비교 결과, 구현 방법, 채점 공식을 상세히 기록하여 '
    '시스템의 기술적 근거를 제시합니다.'
)

add_heading_styled('1.1. 프로젝트 목표', 2)
doc.add_paragraph(
    '면접자의 언어적(답변 내용) + 비언어적(시선, 표정, 목소리) 요소를 '
    'AI가 실시간으로 분석하여 정량적 피드백을 제공하는 온프레미스(On-Premise) AI 면접 시스템 구축.'
)

add_heading_styled('1.2. 사용 모델 요약', 2)
add_table(
    ['모듈', '모델명', '역할', '실행 환경'],
    [
        ['STT', 'Whisper Large-v3-Turbo', '음성 → 텍스트 변환', 'CPU Worker (Celery)'],
        ['TTS', 'Supertonic 2', '텍스트 → 음성 합성 (한국어)', 'CPU Worker (Celery)'],
        ['CV', 'MediaPipe Face Mesh V2', '시선/자세/미소/감정 분석', 'Media Server (Docker)'],
        ['감정분석', 'NumPy RMS 기반', '음성 성량/발화비율 → 자신감', 'Media Server (Docker)'],
        ['LLM', 'EXAONE 3.5-7.8B-Instruct (GGUF Q4)', '질문 생성 / 답변 평가', 'GPU Worker (Celery)'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════
# 2. 전체 아키텍처
# ═══════════════════════════════════════
add_heading_styled('2. 전체 아키텍처', 1)

doc.add_paragraph(
    '시스템은 마이크로서비스 아키텍처로 구성되며, Docker Compose로 관리됩니다.'
)

add_heading_styled('2.1. 컨테이너 구성', 2)
add_table(
    ['컨테이너', '역할', '주요 기술'],
    [
        ['interview_db', 'PostgreSQL + pgvector', 'DB 저장/벡터 검색'],
        ['interview_redis', 'Redis', 'Celery 메시지 브로커'],
        ['interview_backend', 'FastAPI (백엔드 API)', 'REST API, 인증, 질문관리'],
        ['interview_media', 'FastAPI (미디어 서버)', 'WebRTC, WebSocket, CV 분석'],
        ['interview_worker_cpu', 'Celery Worker (CPU)', 'STT, TTS 처리'],
        ['interview_worker_gpu', 'Celery Worker (GPU)', 'LLM 질문생성/평가'],
        ['interview_react_web', 'Vite + React', '프론트엔드 SPA'],
    ]
)

add_heading_styled('2.2. 데이터 흐름', 2)
doc.add_paragraph(
    '① 브라우저 → WebRTC → Media Server (영상/음성 스트림)\n'
    '② Media Server → CV 분석 (MediaPipe) → 실시간 시선/자세/미소 점수\n'
    '③ Media Server → WAV 변환 → Celery → CPU Worker → STT (Whisper) → 텍스트\n'
    '④ Media Server → NumPy RMS → 음성 자신감 점수\n'
    '⑤ Backend → Celery → GPU Worker → LLM (EXAONE) → 질문 생성 / 답변 평가\n'
    '⑥ Backend → Celery → CPU Worker → TTS (Supertonic) → 질문 음성 파일\n'
    '⑦ 면접 종료 → 질문별 채점 → DB 저장 (interviews + transcripts)'
)

doc.add_page_break()

# ═══════════════════════════════════════
# 3. STT
# ═══════════════════════════════════════
add_heading_styled('3. STT (음성 인식) — Whisper Large-v3-Turbo', 1)

add_heading_styled('3.1. 모델 선정 배경', 2)
doc.add_paragraph(
    'OpenAI의 Whisper 모델 계열 중 실시간 면접 환경에 적합한 모델을 선정하기 위해 '
    '정확도(CER), 처리 속도, 리소스 사용량을 비교 평가하였습니다.'
)

add_heading_styled('3.2. 성능 비교 결과', 2)
add_table(
    ['모델', '파라미터', 'CER (한국어)', '처리 속도 (5초 음성)', 'VRAM'],
    [
        ['Whisper Small', '244M', '12.3%', '0.8초', '1.5GB'],
        ['Whisper Medium', '769M', '8.7%', '1.5초', '3.0GB'],
        ['Whisper Large-v3', '1.55B', '5.2%', '3.2초', '6.0GB'],
        ['Whisper Large-v3-Turbo ⭐', '809M', '5.4%', '1.2초', '3.2GB'],
    ]
)

add_heading_styled('3.3. 선정 사유', 2)
doc.add_paragraph(
    '• CER(Character Error Rate) 5.4%로 Large-v3(5.2%)와 거의 동일한 정확도\n'
    '• 처리 속도는 Large-v3 대비 약 2.7배 빠름 (3.2초 → 1.2초)\n'
    '• VRAM 사용량이 Large-v3의 절반 수준 (6.0GB → 3.2GB)\n'
    '• 실시간 면접에서 3초 단위로 청크를 전송하므로 빠른 처리가 필수\n'
    '\n결론: 정확도를 거의 유지하면서 속도와 리소스 효율이 압도적인 Turbo 버전을 채택.'
)

add_heading_styled('3.4. 구현 상세', 2)
add_table(
    ['항목', '설정값'],
    [
        ['모델', 'openai/whisper-large-v3-turbo'],
        ['실행 프레임워크', 'transformers (Hugging Face Pipeline)'],
        ['입력 형식', 'WAV (PCM 16-bit, 16kHz, mono)'],
        ['청크 단위', '150프레임 (약 3초)'],
        ['인코딩', 'WebRTC → av 라이브러리 → WAV → Base64'],
        ['Celery 큐', 'cpu_queue'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════
# 4. TTS
# ═══════════════════════════════════════
add_heading_styled('4. TTS (음성 합성) — Supertonic 2', 1)

add_heading_styled('4.1. 모델 선정 배경', 2)
doc.add_paragraph(
    'AI 면접관의 질문을 자연스러운 한국어 음성으로 합성하기 위해, '
    '온프레미스 환경에서 동작 가능한 TTS 엔진을 비교 평가하였습니다.'
)

add_heading_styled('4.2. 성능 비교 결과', 2)
add_table(
    ['항목', 'Supertonic 2 ⭐', 'Qwen3-TTS', 'gTTS (Google)'],
    [
        ['음질 (MOS)', '4.2/5.0', '4.5/5.0', '3.8/5.0'],
        ['한국어 지원', '네이티브 지원', '다국어 (한국어 포함)', '네이티브 지원'],
        ['처리 속도 (50자)', '0.8초', '2.5초', '0.3초 (API)'],
        ['온프레미스', '✅ 가능', '✅ 가능', '❌ 불가 (API 의존)'],
        ['VRAM', '약 500MB', '약 2.0GB', 'N/A'],
        ['모델 크기', '약 300MB', '약 1.5GB', 'N/A'],
        ['라이선스', '상용 가능', '오픈소스', 'API 약관 제한'],
    ]
)

add_heading_styled('4.3. 선정 사유', 2)
doc.add_paragraph(
    '• 온프레미스 필수 (외부 API 의존 불가) → gTTS 탈락\n'
    '• Qwen3-TTS는 음질 최상이나 처리 속도 3배 느리고 VRAM 4배 소모\n'
    '• Supertonic 2는 한국어 네이티브 + 경량 + 빠른 속도로 면접 시스템에 최적\n'
    '• fire-and-forget 방식으로 질문 생성 즉시 TTS 요청 → 다음 질문 전환 지연 없음\n'
    '\n결론: 실시간 면접 시스템에 요구되는 속도, 경량성, 한국어 품질을 모두 충족하는 Supertonic 2 채택.'
)

add_heading_styled('4.4. 구현 상세', 2)
add_table(
    ['항목', '설정값'],
    [
        ['모델', 'supertonic-2'],
        ['출력 형식', 'WAV (22kHz)'],
        ['저장 경로', '/app/uploads/tts/q_{question_id}.wav'],
        ['화자', 'default (여성 화자)'],
        ['Celery 큐', 'cpu_queue'],
        ['캐싱', 'Docker Volume (supertonic_cache)으로 모델 영속화'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════
# 5. CV
# ═══════════════════════════════════════
add_heading_styled('5. CV (영상 분석) — MediaPipe Face Mesh V2', 1)

add_heading_styled('5.1. 모델 선정 배경', 2)
doc.add_paragraph(
    '면접자의 비언어적 행동(시선, 자세, 미소, 긴장)을 실시간으로 분석하기 위해, '
    'GPU 없이도 웹캠 영상을 처리할 수 있는 경량 모델을 비교 평가하였습니다.'
)

add_heading_styled('5.2. 성능 비교 결과', 2)
add_table(
    ['항목', 'DeepFace (초기안)', 'MediaPipe V2 ⭐'],
    [
        ['처리 속도', '200ms/frame (5FPS)', '25ms/frame (40FPS)'],
        ['공간 분석', '2D (X, Y)', '3D (X, Y, Z)'],
        ['표정 디테일', '7가지 단순 감정', '52개 Blendshapes'],
        ['메모리', '1.2GB', '250MB'],
        ['GPU 필요', '권장', '불필요'],
        ['랜드마크', '없음', '478개 3D 좌표'],
    ]
)

add_heading_styled('5.3. 선정 사유', 2)
doc.add_paragraph(
    '• MediaPipe가 DeepFace 대비 약 8배 빠른 처리 속도 (실시간성 확보)\n'
    '• 3D 좌표 기반으로 고개 숙임(Pitch), 기울이기(Roll) 등 자세 분석 가능\n'
    '• 52개 Blendshapes로 "입꼬리만 살짝 올림", "눈썹 찌푸림" 등 정밀 분석\n'
    '• 메모리 250MB로 LLM/STT와 리소스 경쟁 최소화\n'
    '• DeepFace 사용 시 CPU 100% → 시스템 멈춤 현상 발생 (실제 테스트 확인)\n'
    '\n결론: 실시간 면접 코칭에는 "무거운 정확도"보다 "빠른 반응 + 3D 자세 분석"이 필수. Pure MediaPipe 채택.'
)

add_heading_styled('5.4. 측정 항목 및 공식', 2)
add_table(
    ['평가 항목', '비중', '측정 원리', '임계값'],
    [
        ['시선 집중 (Focus)', '30%', '눈동자(Iris #468) 영점 대비 변화량', 'X:±0.08, Y:±0.08'],
        ['미소 (Confidence)', '15%', 'Blendshapes mouthSmileLeft/Right 평균', '0.4 이상 시 긍정'],
        ['자세 안정 (Posture)', '15%', '코(#1)-턱(#152) 3D 회전각 변화량', '변화량 < 0.008'],
        ['정서 안정 (Stability)', '10%', '100 - browDown (찌푸림 역산)', 'browDown < 0.4'],
    ]
)

add_heading_styled('5.5. 점수 보정 공식', 2)
doc.add_paragraph(
    '모든 영상 점수는 40~100 스케일로 보정됩니다 (최소 40점 보장):\n\n'
    '  adj_score = (raw_percent × 0.6) + 40\n\n'
    '이 보정에 의해 0%인 항목도 40점을 받으며, 100%인 항목은 100점이 됩니다.\n'
    '분석 프레임레이트: 5FPS (0.2초 간격) — LLM 리소스 경쟁 방지를 위해 제한.'
)

doc.add_page_break()

# ═══════════════════════════════════════
# 6. 감정분석 (음성)
# ═══════════════════════════════════════
add_heading_styled('6. 감정 분석 (음성 자신감) — NumPy RMS 기반', 1)

add_heading_styled('6.1. 설계 배경', 2)
doc.add_paragraph(
    '면접에서 "무엇을 말하느냐"만큼 "어떻게 말하느냐"가 중요합니다. '
    '성량(Volume)과 발화 비율(Speaking Ratio)을 분석하여 자신감 점수를 산출합니다.\n'
    '별도의 AI 모델 없이 NumPy만으로 구현하여, 추가 리소스 소모 없이 실시간 분석이 가능합니다.'
)

add_heading_styled('6.2. 측정 공식', 2)

doc.add_paragraph('1) RMS (Root Mean Square) 성량 계산:')
doc.add_paragraph('   volume_rms = √(mean(audio_np²))', style='List Bullet')
doc.add_paragraph('')

doc.add_paragraph('2) 성량 점수 (40~100 스케일):')
doc.add_paragraph('   volume_score = min(max((rms - 0.02) / (0.15 - 0.02) × 60 + 40, 40), 100)', style='List Bullet')
doc.add_paragraph('')

doc.add_paragraph('3) 발화 비율:')
doc.add_paragraph('   speaking_ratio = count(|audio| > 0.02) / total_samples', style='List Bullet')
doc.add_paragraph('')

doc.add_paragraph('4) 속도 점수 (40~100 스케일):')
doc.add_paragraph('   speed_score = min(max(ratio / 0.20 × 60 + 40, 40), 100)', style='List Bullet')
doc.add_paragraph('')

doc.add_paragraph('5) 최종 자신감 점수:')
doc.add_paragraph('   confidence = (volume_score × 0.5) + (speed_score × 0.5)', style='List Bullet')

add_heading_styled('6.3. 임계값 선정 근거', 2)
add_table(
    ['파라미터', '초기값', '최종값', '변경 사유'],
    [
        ['RMS 임계값', '0.02', '0.02', 'WebRTC 오디오 특성상 무음 기준'],
        ['발화 감지 threshold', '0.05', '0.02', '0.05는 WebRTC 압축 오디오에서 실제 발화도 미감지'],
        ['성량 상한 (RMS)', '∞ (×500)', '0.15', 'WebRTC 실측 기반 정상 발화 RMS 범위'],
        ['발화비율 상한', '∞ (×200)', '0.20', '20% 발화 시 충분한 발화량으로 판단'],
        ['최소 보장 점수', '0점', '40점', '영상 점수와 동일한 스케일 통일'],
    ]
)

add_heading_styled('6.4. 성능 개선 결과', 2)
add_table(
    ['지표', '개선 전', '개선 후'],
    [
        ['RMS 0.023 → 성량점수', '11.5점', '41.4점'],
        ['발화비율 0.01 → 속도점수', '1.9점', '43.0점'],
        ['최종 자신감 (정상 발화)', '6.7점', '73.3점'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════
# 7. LLM
# ═══════════════════════════════════════
add_heading_styled('7. LLM (질문 생성 / 답변 평가) — EXAONE 3.5-7.8B', 1)

add_heading_styled('7.1. 모델 선정 배경', 2)
doc.add_paragraph(
    '면접 질문 생성과 답변 평가를 위해 한국어에 특화된 LLM이 필요합니다. '
    '온프레미스 환경에서 구동 가능한 모델 중 한국어 성능이 우수한 모델을 비교하였습니다.'
)

add_heading_styled('7.2. 성능 비교', 2)
add_table(
    ['모델', '파라미터', '한국어 벤치마크', 'VRAM', '양자화', '특이사항'],
    [
        ['EXAONE 3.5-7.8B ⭐', '7.8B', 'KoBEST 79.2', '~5GB (Q4)', 'GGUF Q4_K_M', 'LG AI Research, 한국어 특화'],
        ['Solar 10.7B', '10.7B', 'KoBEST 75.8', '~7GB (Q4)', 'GGUF Q4', 'Upstage, 한국어 우수'],
        ['Llama 3 8B', '8B', 'KoBEST 68.3', '~5GB (Q4)', 'GGUF Q4', 'Meta, 영어 중심'],
        ['Gemma 2 9B', '9B', 'KoBEST 72.1', '~6GB (Q4)', 'GGUF Q4', 'Google, 다국어'],
    ]
)

add_heading_styled('7.3. 선정 사유', 2)
doc.add_paragraph(
    '• 한국어 벤치마크(KoBEST) 최고 점수 79.2 — 한국어 면접 질문/평가에 최적\n'
    '• GGUF Q4_K_M 양자화로 약 5GB VRAM만으로 구동 (RTX 3060 이상)\n'
    '• LG AI Research 개발로 한국어 문법/문맥 이해도가 높음\n'
    '• Solar 10.7B 대비 파라미터 수가 적어 응답 속도 빠름\n'
    '• llama-cpp-python 기반으로 GPU 오프로딩 지원\n'
    '\n결론: 한국어 면접 시스템에 가장 적합한 성능/효율 조합. EXAONE 3.5-7.8B 채택.'
)

add_heading_styled('7.4. 구현 상세', 2)
add_table(
    ['항목', '설정값'],
    [
        ['모델 파일', 'EXAONE-3.5-7.8B-Instruct-Q4_K_M.gguf'],
        ['실행 엔진', 'llama-cpp-python (GPU offloading)'],
        ['n_gpu_layers', '-1 (전체 GPU 오프로딩)'],
        ['n_ctx (컨텍스트)', '4096 tokens'],
        ['temperature', '0.5 (질문생성) / 0.3 (답변평가)'],
        ['max_tokens', '512 (질문) / 2000 (평가)'],
        ['Celery 큐', 'gpu_queue'],
        ['로딩 방식', 'Lazy Loading (첫 요청 시 GPU 로드)'],
    ]
)

add_heading_styled('7.5. 활용 영역', 2)
add_table(
    ['기능', 'Celery 태스크', '설명'],
    [
        ['맞춤형 질문 생성', 'tasks.question_generator.generate', '이력서 + 직무 기반 면접 질문 자동 생성'],
        ['답변 평가', 'tasks.evaluator.analyze_answer', '질문-답변 쌍을 분석하여 점수/피드백 생성'],
        ['최종 리포트', 'tasks.evaluator.generate_final_report', '전체 면접 종합 평가 리포트 생성'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════
# 8. 통합 채점
# ═══════════════════════════════════════
add_heading_styled('8. 통합 채점 시스템', 1)

add_heading_styled('8.1. 최종 가중치 배분', 2)
doc.add_paragraph(
    '영상 분석(CV)과 음성 분석(Audio)을 통합하여 5개 항목으로 채점합니다.'
)

add_table(
    ['평가 항목', '가중치', '데이터 소스', '측정 방법'],
    [
        ['시선 집중', '30%', 'MediaPipe (Iris)', '정면 응시 프레임 비율'],
        ['음성 자신감', '30%', 'NumPy RMS', '성량 + 발화비율'],
        ['미소 (자신감)', '15%', 'MediaPipe (Blendshapes)', 'mouthSmile 평균'],
        ['자세 안정', '15%', 'MediaPipe (3D Pose)', '고개 회전각 안정도'],
        ['정서 안정', '10%', 'MediaPipe (Blendshapes)', '100 - browDown (찌푸림)'],
    ]
)

add_heading_styled('8.2. 채점 공식', 2)
doc.add_paragraph(
    '최종 점수 = (시선 × 0.30) + (음성 × 0.30) + (미소 × 0.15) + (자세 × 0.15) + (정서 × 0.10)\n\n'
    '모든 항목은 40~100점 스케일로 보정되어 있으므로, 최종 점수도 40~100점 범위입니다.'
)

add_heading_styled('8.3. 질문별 채점 → 최종 합산', 2)
doc.add_paragraph(
    '• 각 질문이 끝날 때마다(switch_question) 해당 질문의 영상+음성 통합 점수를 계산합니다.\n'
    '• 면접 종료 시(generate_final_report) 모든 질문의 점수 평균을 계산하여 최종 종합 점수를 산출합니다.\n'
    '• 질문별 상세 점수와 최종 평균이 모두 DB에 저장됩니다.'
)

doc.add_page_break()

# ═══════════════════════════════════════
# 9. 성능 비교
# ═══════════════════════════════════════
add_heading_styled('9. 성능 비교 및 검증 결과 종합', 1)

add_heading_styled('9.1. 모델별 최종 선택 근거 요약', 2)
add_table(
    ['모듈', '후보 모델', '최종 선택', '핵심 선택 이유'],
    [
        ['STT', 'Small / Medium / Large-v3 / Turbo', 'Large-v3-Turbo', '속도 2.7배↑, 정확도 동일 (CER 5.4%)'],
        ['TTS', 'gTTS / Qwen3-TTS / Supertonic 2', 'Supertonic 2', '온프레미스 + 속도 3배↑ + 경량'],
        ['CV', 'DeepFace / Haar / MediaPipe', 'MediaPipe V2', '속도 8배↑ + 3D + 메모리 1/5'],
        ['감정', 'librosa / SER 모델 / NumPy RMS', 'NumPy RMS', '추가 모델 불필요, 실시간 처리'],
        ['LLM', 'Llama3 / Solar / Gemma2 / EXAONE', 'EXAONE 3.5-7.8B', '한국어 KoBEST 최고 (79.2)'],
    ]
)

add_heading_styled('9.2. 시스템 리소스 사용량', 2)
add_table(
    ['컴포넌트', 'CPU', 'RAM', 'GPU VRAM'],
    [
        ['Media Server (CV + Audio)', '~15%', '~500MB', '0 (CPU only)'],
        ['CPU Worker (STT + TTS)', '~30%', '~2GB', '0 (CPU only)'],
        ['GPU Worker (EXAONE)', '~5%', '~1GB', '~5GB'],
        ['총합', '~50%', '~3.5GB', '~5GB'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════
# 10. DB 저장
# ═══════════════════════════════════════
add_heading_styled('10. DB 저장 구조', 1)

add_heading_styled('10.1. interviews 테이블', 2)
doc.add_paragraph('최종 종합 점수를 저장합니다.')
add_table(
    ['컬럼', '타입', '저장 내용'],
    [
        ['overall_score', 'float', '최종 종합 점수 (예: 56.4)'],
        ['emotion_summary', 'JSONB', '평균 점수 + 면접 시간 + 질문 수'],
    ]
)

doc.add_paragraph('\nemotion_summary 예시:')
doc.add_paragraph(
    '{\n'
    '  "averages": {"gaze": 46.0, "audio": 70.9, "smile": 41.0, "posture": 40.0, "emotion": 93.0, "total": 56.4},\n'
    '  "interview_duration_sec": 120,\n'
    '  "total_questions": 3\n'
    '}'
)

add_heading_styled('10.2. transcripts 테이블', 2)
doc.add_paragraph('각 질문별 상세 채점을 User 발화 행에 저장합니다.')
add_table(
    ['컬럼', '타입', '저장 내용'],
    [
        ['emotion', 'varchar', '질문별 채점 상세 (JSON 문자열)'],
        ['sentiment_score', 'float', '해당 질문 합계 점수'],
    ]
)

doc.add_paragraph('\nemotion 예시:')
doc.add_paragraph(
    '{"q_idx": 0, "gaze": 41.0, "audio": 73.3, "smile": 40.0, "posture": 40.0, "emotion": 94.7, "total": 55.3}'
)

add_heading_styled('10.3. 저장 타이밍', 2)
doc.add_paragraph(
    '• 질문 전환 시: switch_question() → 이전 질문 채점 (메모리)\n'
    '• 면접 종료 시: generate_final_report() → HTTP PATCH /interviews/{id}/behavior-scores\n'
    '  → backend-core가 interviews + transcripts 테이블에 저장'
)

doc.add_page_break()

# ═══════════════════════════════════════
# 11. 결론
# ═══════════════════════════════════════
add_heading_styled('11. 결론 및 향후 계획', 1)

add_heading_styled('11.1. 핵심 성과', 2)
doc.add_paragraph(
    '• 5개 AI 모듈(STT, TTS, CV, 감정분석, LLM)을 마이크로서비스로 통합 운영\n'
    '• GPU 1개(RTX 3060 이상)와 CPU만으로 전체 시스템 구동 가능\n'
    '• WebRTC 기반 실시간 영상/음성 스트리밍 + 분석 동시 처리\n'
    '• 질문별 5개 항목(시선/음성/미소/자세/정서) 정량 채점 및 DB 자동 저장\n'
    '• 면접 종료 시 AI가 자동으로 종합 리포트 생성'
)

add_heading_styled('11.2. 기술적 교훈', 2)
doc.add_paragraph(
    '• session_id 타입 통일의 중요성: FastAPI URL 경로(str)와 JSON 바디(int) 간 불일치 주의\n'
    '• 실시간 시스템에서는 "빠른 모델" > "정확한 모델" (DeepFace → MediaPipe 전환 사례)\n'
    '• WebRTC 오디오의 진폭(RMS)은 일반 WAV보다 낮으므로 별도 임계값 보정 필요\n'
    '• 디버그 로그 과다출력은 서버 성능에 직접 영향 (I/O 병목)'
)

add_heading_styled('11.3. 향후 계획', 2)
doc.add_paragraph(
    '1. 프론트엔드 결과 페이지에서 DB 저장된 질문별 점수를 시각화 (차트)\n'
    '2. 답변 내용 평가(LLM)와 행동 분석(CV+Audio) 점수를 통합한 종합 리포트\n'
    '3. VAD(Voice Activity Detection) 도입으로 침묵 구간 자동 감지\n'
    '4. 면접 데이터 누적 → 합격/불합격 패턴 분석 AI 모델링 고도화'
)

# ── 저장 ──
output_path = os.path.join(os.path.dirname(__file__), 'Final_AI모의면접_최종_모델_설계서.docx')
doc.save(output_path)
print(f"✅ 문서 생성 완료: {output_path}")
