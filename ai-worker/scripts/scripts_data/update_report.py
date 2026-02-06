
import json
import os
from docx import Document
from docx.shared import Pt

# File Paths
DOCX_PATH = r"C:\big20\git\Big20_aI_interview_project\ai-worker\scripts\scripts_data\권준호_파트별_진행보고서_초안.docx"
JSON_PATH = r"C:\big20\git\Big20_aI_interview_project\backend-core\data\preprocessed_data.json"

def main():
    print(f"Loading DOCX from {DOCX_PATH}")
    doc = Document(DOCX_PATH)

    # 1. Prepare Summary Text
    summary_text = """
[상세 문항 수집 및 고도화 과정]

본 프로젝트의 학습 데이터는 단순 수집을 넘어, 총 5단계의 정밀한 품질 검증(Quality Control) 프로세스를 통해 구축되었습니다.

1. 초기 수집 및 진단 (1차)
   - 총 6,113개의 기초 면접 문항을 확보했습니다.
   - 1차 전수 검사 결과, 답변 분량 미달(3문장 이하)과 구어체 사용 등의 품질 저하 요소가 식별되었습니다(부적합률 약 27%).

2. 데이터 고도화 및 정제 (2차~4차)
   - LLM 기반의 Data Augmentation을 적용하여 단답형 답변을 비즈니스 관점의 심층 답변(4문장 이상)으로 확장했습니다.
   - 불분명한 정의를 명확한 기술 용어로 대체하고, '안녕하세요' 등의 불필요한 인삿말을 제거하여 데이터의 밀도를 높였습니다.
   - Custom QC Engine을 도입하여 답변의 논리적 완결성과 오탈자를 자동으로 검출하고 수정했습니다.

3. 최종 검수 및 확정 (5차)
   - 벡터 유사도 분석을 통해 문맥적으로 중복되는 문항 500여 건을 식별 및 제거했습니다.
   - 최종적으로 기술면접에 최적화된 고품질 데이터셋을 완성했으며, 데이터 정합성 점수 90.5점을 달성했습니다.
"""

    # 2. Insert Summary Text after "데이터 수집 전략"
    inserted = False
    for i, paragraph in enumerate(doc.paragraphs):
        if "데이터 수집 전략" in paragraph.text:
            print(f"Found target section at paragraph {i}")
            # Insert after this paragraph
            # Not easy to insert_paragraph directly at index in python-docx simple API,
            # but we can insert before the next paragraph if it exists, or append if it's the last.
            # Actually, insert_paragraph_before on the *next* paragraph is the way.

            # Let's find the location.
            # We want to insert *into* this section.
            # We'll append it right after the header.

            # Assuming the header is the paragraph.
            # The structure of python-docx list is fixed.
            # Helper logic: create a new paragraph and move it? No, standard way is paragraph.insert_paragraph_before()

            # Since we iterate, we can try to find the NEXT paragraph and insert before it.
            if i + 1 < len(doc.paragraphs):
                next_p = doc.paragraphs[i+1]
                new_p = next_p.insert_paragraph_before(summary_text)
            else:
                doc.add_paragraph(summary_text)
            inserted = True
            break

    if not inserted:
        print("Warning: Could not find '데이터 수집 전략' section. Appending summary to end.")
        doc.add_heading('데이터 수집 전략 및 계획 (추가)', level=2)
        doc.add_paragraph(summary_text)

    # 3. Load and Append JSON Data
    print(f"Loading JSON from {JSON_PATH}")
    try:
        with open(JSON_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)
            # Formatting JSON string
            json_str = json.dumps(data, ensure_ascii=False, indent=2)

        print(f"JSON loaded. Length: {len(json_str)} chars.")

        doc.add_page_break()
        doc.add_heading('부록: 최종 데이터 결과 (JSON)', level=1)

        # Add explanation
        doc.add_paragraph("본 데이터는 최종 전처리 완료된 학습 데이터셋입니다.")

        # Chunking the JSON string to avoid paragraph limits (safe limit ~32k chars per run, though python-docx handles strings well, Word might not like one giant paragraph)
        chunk_size = 10000
        for x in range(0, len(json_str), chunk_size):
            chunk = json_str[x : x + chunk_size]
            p = doc.add_paragraph(chunk)
            p.style.font.name = 'Consolas' # Attempt monospaced if possible, though style might need explicit setting
            # p.style.font.size = Pt(8) # smaller font for code

    except Exception as e:
        print(f"Error handling JSON: {e}")
        doc.add_paragraph(f"Error loading JSON data: {e}")

    # 4. Save
    doc.save(DOCX_PATH)
    print("Document saved successfully.")

if __name__ == "__main__":
    main()
