
import json
import os
import time
from docx import Document
from docx.shared import Pt

# Constants
DOCX_PATH = r"C:\big20\git\Big20_aI_interview_project\ai-worker\scripts\scripts_data\권준호_파트별_진행보고서_초안.docx"
JSON_PATH = r"C:\big20\git\Big20_aI_interview_project\backend-core\data\preprocessed_data.json"

def main():
    print(f"Processing {DOCX_PATH}...")
    try:
        doc = Document(DOCX_PATH)
    except Exception as e:
        print(f"Error opening document: {e}")
        return

    # --- 1. CLEANUP (Remove JSON Dump & Old Strategy) ---
    paragraphs_to_remove = []

    # Identify Appendix Section (JSON Dump)
    dump_started = False
    for p in doc.paragraphs:
        if "부록: 최종 데이터 결과" in p.text or "부록: 학습 데이터셋 명세" in p.text:
            dump_started = True

        if dump_started:
            paragraphs_to_remove.append(p)

        # Identify Previous Strategy Text to replace
        if "[상세 문항 수집 및 고도화 과정]" in p.text or "[상세 문항 수집 및 품질 고도화 프로세스]" in p.text:
            paragraphs_to_remove.append(p)
        if "단순 수집을 넘어" in p.text or "단순 수집을 지양하고" in p.text:
             paragraphs_to_remove.append(p)

    # Remove them
    print(f"Removing {len(paragraphs_to_remove)} paragraphs of old content...")
    for p in paragraphs_to_remove:
        try:
            p_element = p._element
            if p_element.getparent() is not None:
                p_element.getparent().remove(p_element)
        except:
            pass

    # --- 2. INSERT NEW STRATEGY (CLAUDE.md Based) ---
    new_strategy_text = """
[하이브리드 데이터 구축 프로세스: Human-AI Collaboration]

본 프로젝트는 12명의 참여자가 수집한 원천 데이터를 AI 전문가 모델이 정밀 검증하고 보정하는 'Human-AI 협업 파이프라인'을 통해 구축되었습니다.

(1) 원천 데이터 수집 (Human-Driven Collection)
   - 참여자: 총 12명 (준호, 린, 재민, 여림, 지훈, 승재, 하량, 다윗, 명혁, 윤재, 시현, 유라)
   - 규모: 인당 약 510문항, 총 6,113건의 기술 면접(IT 전반) 원천 데이터 확보.
   - 방식: 12개 채널을 통한 병렬 수집 후 표준화된 JSON 포맷으로 통합.

(2) 전문가 AI 정밀 평가 (Phase 1: Deep Evaluation)
   - 평가 AI: 기술 면접관 페르소나를 탑재한 AI가 F1(의도), F4(사실오류), F5(분량) 등 5대 필수 항목 기준 전수 검사.
   - 결과 진단: 초기 데이터의 약 27%가 '답변 깊이 부족(3문장 이하)' 및 '기술적 모호함'으로 인해 FAIL 판정.

(3) 데이터 고도화 및 품질 보정 (Phase 2: Iterative Refinement)
   - 심층 고도화: FAIL 답변에 대해 '원리-비즈니스 임팩트-제약조건'을 보강하여 5~6문장 분량의 전문가급 답변으로 확장 (1,534건).
   - 오류 교정: 식별된 사실 오류(43건) 및 오탈자 전수 교정, 참여자 고유 스타일을 유지하되 전문성 강화.
   - 품질 재검증: 1차 자동 검수(규격 확인) → 2차 정밀 재평가(등급 확인)의 이중 QC 수행.

(4) 최종 확정 (Finalization)
   - 성과: 12명 전원 데이터 품질 등급 '종합 A등급' 달성.
   - 최종 데이터: 총 6,113건의 무결점(Zero-Defect) 기술 면접 데이터셋 완성.
"""

    # Find insertion point
    target_p = None
    for i, p in enumerate(doc.paragraphs):
        if "데이터 수집 전략" in p.text:
            target_p = p
            break

    if target_p:
        # Insert after header
        # Find index
        try:
            idx = doc.paragraphs.index(target_p)
            if idx + 1 < len(doc.paragraphs):
                doc.paragraphs[idx+1].insert_paragraph_before(new_strategy_text)
            else:
                doc.add_paragraph(new_strategy_text)
        except:
            doc.add_paragraph(new_strategy_text)
    else:
        doc.add_heading("데이터 수집 전략 및 계획 (보강)", level=2)
        doc.add_paragraph(new_strategy_text)


    # --- 3. ADD APPENDIX (Summary Stats & Samples) ---
    print(f"Loading JSON from {JSON_PATH}")
    with open(JSON_PATH, 'r', encoding='utf-8') as f:
        data = json.load(f)

    total_count = len(data)
    # Recalculate stats for accuracy
    cat_counts = {}
    for item in data:
        c = item.get('QuestionCategory', 'Unknown')
        cat_counts[c] = cat_counts.get(c, 0) + 1

    doc.add_page_break()
    doc.add_heading('부록: 학습 데이터셋 명세 (Data Specification)', level=1)

    # 3.1 Statistics
    doc.add_heading('1. 데이터 통계 요약', level=2)
    stat_text = f"총 데이터 개수: {total_count:,}건\n\n카테고리별 분포 (QuestionCategory):"
    for cat, cnt in sorted(cat_counts.items(), key=lambda x: x[1], reverse=True):
        pct = (cnt / total_count) * 100
        stat_text += f"\n  - {cat}: {cnt:,}건 ({pct:.1f}%)"
    doc.add_paragraph(stat_text)

    # 3.2 Schema
    doc.add_heading('2. 데이터 구조', level=2)
    doc.add_paragraph("""- question: 면접 질문
- answer: 전문가 검증 및 확장이 완료된 모범 답변
- QuestionCategory: 평가 영역 (TECHNICAL, BEHAVIORAL 등)
- QuestionDifficulty: 난이도 (EASY, MEDIUM, HARD)""")

    # 3.3 Samples
    doc.add_heading('3. 최종 데이터 샘플 (예시)', level=2)
    # Select sample: one technical, one behavioral
    samples_to_show = []
    for item in data:
        if item.get('QuestionCategory') == 'TECHNICAL':
            samples_to_show.append(item)
            break
    for item in data:
        if item.get('QuestionCategory') == 'BEHAVIORAL':
            samples_to_show.append(item)
            break

    for i, s in enumerate(samples_to_show):
        p = doc.add_paragraph()
        run = p.add_run(f"[Sample #{i+1}] {s['QuestionCategory']} / {s['QuestionDifficulty']}")
        run.bold = True

        # Pretty print subset
        subset = {k: s[k] for k in ['question', 'answer']}
        json_disp = json.dumps(subset, ensure_ascii=False, indent=2)
        p2 = doc.add_paragraph(json_disp)
        p2.style.font.name = "Consolas"
        p2.style.font.size = Pt(9)
        p2.paragraph_format.left_indent = Pt(20)

    # --- 4. SAVE ---
    # Try different filenames if permission denied
    save_candidates = [
        DOCX_PATH,
        DOCX_PATH.replace(".docx", "_revised.docx"),
        DOCX_PATH.replace(".docx", "_v2.docx")
    ]

    saved = False
    for path in save_candidates:
        try:
            doc.save(path)
            print(f"Successfully saved to: {path}")
            saved = True
            break
        except PermissionError:
            print(f"Permission denied for {path}, trying next...")
            time.sleep(1)

    if not saved:
        print("Failed to save document. Please close the file and try again.")

if __name__ == "__main__":
    main()
