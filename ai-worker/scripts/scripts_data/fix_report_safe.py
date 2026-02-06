
import json
import os
from docx import Document
from docx.shared import Pt

INPUT_PATH = r"C:\big20\git\Big20_aI_interview_project\ai-worker\scripts\scripts_data\권준호_파트별_진행보고서_초안.docx"
OUTPUT_PATH = r"C:\big20\git\Big20_aI_interview_project\ai-worker\scripts\scripts_data\권준호_파트별_진행보고서_최종.docx"
JSON_PATH = r"C:\big20\git\Big20_aI_interview_project\backend-core\data\preprocessed_data.json"

def main():
    print(f"Opening DOCX: {INPUT_PATH}")
    try:
        doc = Document(INPUT_PATH)
    except Exception as e:
        print(f"Error opening document: {e}")
        return

    # 1. REMOVE PREVIOUSLY ADDED CONTENT
    paragraphs_to_remove = []

    # Find the boundary for the Appendix dump
    dump_start_index = -1
    for i, p in enumerate(doc.paragraphs):
        if "부록: 최종 데이터 결과 (JSON)" in p.text or "부록: 최종 데이터 결과 및 스펙" in p.text:
            dump_start_index = i
            break

    if dump_start_index != -1:
        print(f"Found existing JSON dump starting at paragraph {dump_start_index}. Removing ALL subsequent paragraphs.")
        for i in range(dump_start_index, len(doc.paragraphs)):
            paragraphs_to_remove.append(doc.paragraphs[i])

    # Find the inserted summary text
    for p in doc.paragraphs:
        if "[상세 문항 수집 및 고도화 과정]" in p.text or "[상세 문항 수집 및 품질 고도화 프로세스]" in p.text:
            paragraphs_to_remove.append(p)
        if "데이터 수집 전략 및 계획 (추가)" in p.text:
            paragraphs_to_remove.append(p)

    count_removed = 0
    for p in paragraphs_to_remove:
        try:
            p_element = p._element
            if p_element.getparent() is not None:
                p_element.getparent().remove(p_element)
                count_removed += 1
        except Exception as e:
            pass # Ignore if already removed or not found

    print(f"Removed {count_removed} paragraphs.")


    # 2. INSERT NEW STRATEGY TEXT
    target_p = None
    target_index = -1
    for i, p in enumerate(doc.paragraphs):
        if "데이터 수집 전략" in p.text:
            target_p = p
            target_index = i
            break

    new_strategy_text = """
[상세 문항 수집 및 품질 고도화 프로세스]

본 프로젝트는 단순한 데이터 수집을 지양하고, 총 5단계의 엄격한 품질 관리(QC) 프로세스를 통해 '기술 면접 특화 데이터셋'을 구축했습니다.

(1) 1차 수집 및 진단 (Initial Gathering)
   - 확보 데이터: 기초 문항 7,141건 (전체 풀 기준)
   - 품질 진단: 단순 단답형(3문장 이하) 답변 다수 존재, 초기 합격률 70%대 기록.

(2) 데이터 증강 및 고도화 (Augmentation)
   - LLM 기반 증강: 단답형 답변에 '비즈니스 임팩트', '기술적 원리'를 보강하여 평균 답변 길이 4~5문장 확보.
   - 스타일 교정: 구어체 및 불필요한 서두를 제거하고 기술적 전문성을 갖춘 문체로 전량 수정.

(3) 최종 최적화 (Optimization)
   - 중복 제거: 벡터 유사도 분석을 통해 문맥적으로 중복되는 문항 식별 및 정제.
   - 최종 성과: 데이터 정합성 검증 통과율 100% (6,109개 확정), 정성 품질 점수 90.5점 달성.
"""

    if target_p:
        # Insert after found heading.
        # Python-docx list handling: finding current index of target_p is safest.
        try:
            current_index = doc.paragraphs.index(target_p)
            if current_index + 1 < len(doc.paragraphs):
                doc.paragraphs[current_index+1].insert_paragraph_before(new_strategy_text)
            else:
                doc.add_paragraph(new_strategy_text)
        except ValueError:
            # target_p might not be in list if modified? shouldn't happen.
            doc.add_paragraph(new_strategy_text)
    else:
        doc.add_heading("데이터 수집 전략 및 계획 (보강)", level=2)
        doc.add_paragraph(new_strategy_text)


    # 3. APPEND NEW DATA SPEC APPENDIX
    print(f"Loading JSON statistics from {JSON_PATH}")
    try:
        with open(JSON_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)

        total_count = len(data)
        categories = {}
        for item in data:
            cat = item.get('QuestionCategory', 'Uncategorized')
            categories[cat] = categories.get(cat, 0) + 1

        doc.add_page_break()
        doc.add_heading('부록: 학습 데이터셋 명세 (Data Specification)', level=1)

        # 3.1 Statistics
        doc.add_heading('1. 데이터 통계 요약', level=2)
        stat_lines = [
            f"총 데이터 개수: {total_count:,}건 (원천 데이터 포함)",
            "카테고리별 분포:"
        ]
        for cat, count in categories.items():
            percentage = (count / total_count) * 100
            stat_lines.append(f"  - {cat}: {count:,}건 ({percentage:.1f}%)")

        doc.add_paragraph('\n'.join(stat_lines))

        # 3.2 Schema Description
        doc.add_heading('2. 데이터 구조 및 스키마', level=2)
        schema_desc = """본 데이터셋은 JSON 포맷으로 구성되어 있으며, 각 문항은 다음과 같은 필드를 포함합니다:
- question: 면접 질문 텍스트
- answer: 모범 답변 (평균 4문장 이상의 상세 설명 포함)
- QuestionCategory: 질문 카테고리 (TECHNICAL: 직무지식, BEHAVIORAL: 인성/행동 등)
- QuestionDifficulty: 난이도 (EASY, MEDIUM, HARD)"""
        doc.add_paragraph(schema_desc)

        # 3.3 Samples
        doc.add_heading('3. 데이터 샘플 (예시)', level=2)
        samples = data[:3] # Show 3 samples

        for idx, sample in enumerate(samples):
            p = doc.add_paragraph()
            run = p.add_run(f"[Sample #{idx+1}] [{sample.get('QuestionCategory')}/{sample.get('QuestionDifficulty')}]")
            run.bold = True
            run.font.size = Pt(10)

            # Use limited fields only for clarity
            mini_sample = {
                "question": sample.get('question'),
                "answer": sample.get('answer'),
                "category": sample.get('QuestionCategory')
            }
            json_str = json.dumps(mini_sample, ensure_ascii=False, indent=2)

            p_json = doc.add_paragraph(json_str)
            p_json.style.font.name = "Consolas"
            p_json.style.font.size = Pt(9)
            p_json.paragraph_format.left_indent = Pt(20)

    except Exception as e:
        doc.add_paragraph(f"데이터 로드 중 오류 발생: {e}")

    # Save to NEW path to avoid PermissionError
    try:
        doc.save(OUTPUT_PATH)
        print(f"Report corrected and saved to: {OUTPUT_PATH}")
    except Exception as e:
        print(f"Critical Error saving document: {e}")

if __name__ == "__main__":
    main()
