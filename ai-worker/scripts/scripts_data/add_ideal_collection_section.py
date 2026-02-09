
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os

# 파일 경로
DOCX_PATH = r"C:\big20\git\Big20_aI_interview_project\ai-worker\scripts\scripts_data\권준호_파트별_진행보고서_찐최종.docx"
CORP_DATA_PATH = r"C:\big20\git\Big20_aI_interview_project\backend-core\data\corp_data.json"
OUTPUT_PATH = r"C:\big20\git\Big20_aI_interview_project\ai-worker\scripts\scripts_data\권준호_파트별_진행보고서_인재상추가.docx"

def count_ideal_data():
    """corp_data.json에서 인재상 데이터 통계 계산"""
    with open(CORP_DATA_PATH, 'r', encoding='utf-8') as f:
        corp_data = json.load(f)

    total_count = len(corp_data)
    collected_count = sum(1 for item in corp_data if item.get('ideal', '').strip())
    empty_count = total_count - collected_count

    return {
        'total': total_count,
        'collected': collected_count,
        'empty': empty_count,
        'collection_rate': (collected_count / total_count * 100) if total_count > 0 else 0
    }

def add_ideal_collection_section(doc, stats):
    """인재상 데이터 수집 섹션 추가"""

    # 새로운 섹션 추가
    doc.add_heading('2.4. 기업 인재상 데이터 수집 (Manual Collection)', level=2)

    # 개요 단락
    overview = doc.add_paragraph()
    overview.add_run('[개요] ').bold = True
    overview.add_run(
        f'총 {stats["total"]}개 기업의 인재상 데이터 중 {stats["collected"]}개를 수동으로 수집 완료하였으며, '
        f'수집률 {stats["collection_rate"]:.1f}%를 달성하였습니다.'
    )

    # 공수 및 자동화 분석
    doc.add_paragraph()
    effort_heading = doc.add_paragraph()
    effort_heading.add_run('(1) 공수 분석 및 자동화 필요성').bold = True

    # 공수 분석 표
    effort_para = doc.add_paragraph()
    effort_para.add_run('[수동 수집 공수]\n').bold = True
    effort_content = f"""- 총 작업 인원: 12명
- 수집 완료 건수: {stats["collected"]}개 기업
- 인당 평균 수집 건수: {stats["collected"] // 12}개 (약 {stats["collected"] / 12:.1f}개)
- 인당 소요 시간: 약 1시간
- 총 투입 공수: 12인·시간

[미수집 데이터]
- 미수집 건수: {stats["empty"]}개 기업
- 예상 추가 공수: 약 {stats["empty"] / 12:.1f}시간 (12명 기준)
"""
    effort_para.add_run(effort_content)

    # 자동화 필요성
    doc.add_paragraph()
    auto_heading = doc.add_paragraph()
    auto_heading.add_run('(2) 자동화 전략 및 개선 방향').bold = True

    auto_content = doc.add_paragraph()
    auto_content.add_run('[현황 및 문제점]\n').bold = True
    auto_content.add_run(
        '기업 인재상 데이터는 각 기업의 채용 페이지, IR 자료, 공식 홈페이지 등에 산재되어 있어 '
        '수동 수집 시 다음과 같은 한계가 존재합니다:\n'
        '  - 시간 소모: 기업당 평균 5~7분 소요 (검색 → 확인 → 복사 → 정리)\n'
        '  - 일관성 부족: 수집자마다 표현 방식과 상세도가 상이\n'
        '  - 업데이트 지연: 기업의 인재상 변경 시 실시간 반영 불가\n'
        '  - 확장성 제약: 신규 기업 추가 시 반복적인 수동 작업 필요\n'
    )

    doc.add_paragraph()
    auto_solution = doc.add_paragraph()
    auto_solution.add_run('[자동화 방안]\n').bold = True
    auto_solution.add_run(
        '웹 크롤링 및 AI 기반 자동 수집 시스템 구축을 통해 다음과 같은 개선이 가능합니다:\n\n'
        '① 웹 크롤링 자동화\n'
        '  - Selenium/BeautifulSoup 기반 채용 페이지 자동 탐색\n'
        '  - 주요 채용 플랫폼 (사람인, 잡코리아, 링크드인) API 연동\n'
        '  - 정기적 스케줄링을 통한 데이터 업데이트 (주 1회)\n\n'
        '② AI 기반 텍스트 추출 및 정제\n'
        '  - LLM을 활용한 인재상 키워드 자동 추출\n'
        '  - 다양한 표현을 표준화된 형식으로 변환\n'
        '  - 중복 및 불필요한 내용 자동 필터링\n\n'
        '③ 품질 검증 프로세스\n'
        '  - 자동 수집 데이터의 신뢰도 점수 산출\n'
        '  - 낮은 신뢰도 데이터는 수동 검수 대상으로 분류\n'
        '  - 최종 데이터는 Human-in-the-Loop 방식으로 검증\n'
    )

    # 기대 효과
    doc.add_paragraph()
    effect_heading = doc.add_paragraph()
    effect_heading.add_run('(3) 자동화 기대 효과').bold = True

    effect_content = doc.add_paragraph()
    effect_content.add_run(
        f'자동화 시스템 구축 시 다음과 같은 효과를 기대할 수 있습니다:\n\n'
        f'[효율성 개선]\n'
        f'  - 현재 미수집 {stats["empty"]}개 기업 데이터를 자동으로 수집 가능\n'
        f'  - 수동 작업 대비 약 90% 시간 절감 (12인·시간 → 1.2인·시간)\n'
        f'  - 신규 기업 추가 시 즉시 자동 수집 가능\n\n'
        f'[품질 향상]\n'
        f'  - 표준화된 형식으로 일관성 확보\n'
        f'  - 실시간 업데이트를 통한 최신성 유지\n'
        f'  - AI 검증을 통한 정확도 향상\n\n'
        f'[확장성 확보]\n'
        f'  - 현재 {stats["total"]}개 기업에서 1,000개 이상 기업으로 확장 가능\n'
        f'  - 다양한 산업군 및 직무별 인재상 데이터 수집 가능\n'
        f'  - 글로벌 기업 인재상 데이터 수집으로 확장 가능\n'
    )

    # 현재 수집 데이터 샘플
    doc.add_paragraph()
    sample_heading = doc.add_paragraph()
    sample_heading.add_run('(4) 수집 데이터 샘플').bold = True

    # corp_data.json에서 샘플 추출
    with open(CORP_DATA_PATH, 'r', encoding='utf-8') as f:
        corp_data = json.load(f)

    # ideal이 있는 데이터 중 3개 샘플 추출
    samples = [item for item in corp_data if item.get('ideal', '').strip()][:3]

    for i, sample in enumerate(samples, 1):
        sample_para = doc.add_paragraph()
        sample_para.add_run(f'[샘플 #{i}] {sample["name"]}\n').bold = True
        sample_para.add_run(f'인재상: {sample["ideal"]}\n')
        if sample.get('description'):
            sample_para.add_run(f'기업 설명: {sample["description"][:100]}...\n')
        sample_para.paragraph_format.left_indent = Pt(20)

    return doc

def main():
    print(f"Processing {DOCX_PATH}...")

    # 문서 로드
    try:
        doc = Document(DOCX_PATH)
        print("✓ Document loaded successfully")
    except Exception as e:
        print(f"✗ Error loading document: {e}")
        return

    # 인재상 데이터 통계 계산
    print("\nCalculating ideal data statistics...")
    stats = count_ideal_data()
    print(f"✓ Total companies: {stats['total']}")
    print(f"✓ Collected: {stats['collected']} ({stats['collection_rate']:.1f}%)")
    print(f"✓ Empty: {stats['empty']}")

    # 인재상 수집 섹션 추가
    print("\nAdding ideal collection section...")
    doc = add_ideal_collection_section(doc, stats)
    print("✓ Section added successfully")

    # 문서 저장
    print(f"\nSaving document to {OUTPUT_PATH}...")
    try:
        doc.save(OUTPUT_PATH)
        print(f"✓ Document saved successfully: {OUTPUT_PATH}")
    except Exception as e:
        print(f"✗ Error saving document: {e}")
        # 대체 경로로 저장 시도
        alt_path = OUTPUT_PATH.replace('.docx', '_backup.docx')
        try:
            doc.save(alt_path)
            print(f"✓ Document saved to alternative path: {alt_path}")
        except Exception as e2:
            print(f"✗ Failed to save to alternative path: {e2}")

if __name__ == "__main__":
    main()
