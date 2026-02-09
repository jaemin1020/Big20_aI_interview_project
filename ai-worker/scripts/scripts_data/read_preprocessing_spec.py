from docx import Document
import os

# 전처리 명세서 읽기
spec_path = r"C:\big20\git\Big20_aI_interview_project\ai-worker\scripts\scripts_data\데이터 전처리 명세서 초안.docx"

print("=== 데이터 전처리 명세서 내용 ===\n")

try:
    doc = Document(spec_path)

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            # 스타일 정보 포함
            style = para.style.name if para.style else "Normal"
            print(f"[{i}] [{style}] {para.text}")

    print(f"\n총 {len(doc.paragraphs)}개 단락")

    # 테이블이 있으면 출력
    if doc.tables:
        print(f"\n=== 테이블 {len(doc.tables)}개 발견 ===")
        for t_idx, table in enumerate(doc.tables):
            print(f"\n[Table {t_idx + 1}]")
            for r_idx, row in enumerate(table.rows):
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                print(f"  Row {r_idx}: {row_text}")

except Exception as e:
    print(f"Error: {e}")
