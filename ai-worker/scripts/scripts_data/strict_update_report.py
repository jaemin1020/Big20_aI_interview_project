
import json
import os
import time
from docx import Document

# Constants
DOCX_PATH = r"C:\big20\git\Big20_aI_interview_project\ai-worker\scripts\scripts_data\권준호_파트별_진행보고서_초안.docx"

def main():
    print(f"Strictly refining {DOCX_PATH} based on CLAUDE.md only...")
    try:
        doc = Document(DOCX_PATH)
    except Exception as e:
        print(f"Error opening document: {e}")
        return

    # --- 1. REMOVE PREVIOUS SECTION ---
    paragraphs_to_remove = []
    markers = [
        "하이브리드 데이터 파이프라인",
        "Tech Lead 페르소나",
        "벡터 유사도",
        "중복 제거",
        "상세 문항 수집 및 고도화",
        "Human-AI 협업 파이프라인"
    ]

    for p in doc.paragraphs:
        if any(m in p.text for m in markers) or "전문가 모델이 정밀 검증하고" in p.text:
            paragraphs_to_remove.append(p)
        if "원천 데이터 수집 (Human-" in p.text or "원천 데이터 확보 (Large-" in p.text:
             paragraphs_to_remove.append(p)

    for p in paragraphs_to_remove:
        try:
            p._element.getparent().remove(p._element)
        except:
            pass

    # --- 2. INSERT STRICT TRUTH CONTENT ---
    # Based STRICTLY on C:\big20\git\Big20_aI_interview_project\ai-worker\scripts\scripts_data\CLAUDE.md

    strict_strategy_text = """
[데이터 수집 및 품질 고도화 전략]

본 프로젝트는 12명의 참여자가 수집한 원천 데이터를 IT 전문가 관점에서 전수 검사하고, 결함이 발견된 문항을 집중 개선하는 'Targeted Refinement' 방식으로 수행되었습니다.

1. 원천 데이터 수집 (Source Collection)
   - 참여자: 12명 (준호, 린, 재민 등 파트별 담당자 전원)
   - 방식: 인당 11개 Chunk(약 510문항)로 분할하여 병렬 수집 수행.
   - 확보 규모: IT 기술 면접 전반(OS, NW, DB 등) 문항 총 6,113건 확보 (CLAUDE.md 기준).

2. 전문가 관점의 정밀 검증 (Expert Evaluation)
   - 검증 주체: IT 전반 분야 전문가(IT Generalist Expert) 기준을 적용.
   - 5대 필수 통과 조건(FAIL Criteria) 적용:
     (F1) 질문 의도 불일치, (F4) 사실 오류, (F5) 4문장 이하 등 엄격한 기준 적용.
   - 진단 결과: 전체 6,113건 중 1,534건(약 25%)에 대해 개선 필요(FAIL) 판정.

3. 데이터 고도화 실행 (Quality Improvement Phase)
   - 확장(Expansion): 4문장 이하의 단답형 답변 1,534건을 '트레이드오프, 실무 사례, 제약 조건'을 포함한 5~6문장 분량으로 확장.
   - 교정(Correction): 검증된 사실 오류 43건 및 오탈자 전수 교정.
   - 원칙: 원본 데이터를 보존하되, IT 학술적 정확성과 인과관계를 보강하는 방향으로 수정.

4. 최종 품질 확정 (Finalization)
   - 이중 검수(Double Check): '1차 자동 검증(규격)' 및 '2차 등급 재평가(S~D등급)' 수행.
   - 최종 성과: 12명 전원의 데이터가 종합 평가 등급 **'A(우수) 이상'**을 달성하며 프로젝트 목표 충족.
"""

    # --- 3. INSERTION ---
    target_p = None
    for i, p in enumerate(doc.paragraphs):
        if "데이터 수집 전략" in p.text:
            target_p = p
            break

    if target_p:
        try:
            idx = doc.paragraphs.index(target_p)
            if idx + 1 < len(doc.paragraphs):
                doc.paragraphs[idx+1].insert_paragraph_before(strict_strategy_text)
            else:
                doc.add_paragraph(strict_strategy_text)
        except:
             doc.add_paragraph(strict_strategy_text)
    else:
        doc.add_heading("데이터 수집 전략 및 계획", level=2)
        doc.add_paragraph(strict_strategy_text)

    # Save
    try:
        doc.save(DOCX_PATH)
        print("Document updated with strict CLAUDE.md facts.")
    except Exception as e:
        print(f"Save failed: {e}")

if __name__ == "__main__":
    main()
