
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(base_path, output_filename):
    files = [
        "01-파싱.md",
        "02-청킹.md",
        "03.엑사원모델.md",
        "04.임베딩.md",
        "05.pgvector.md",
        "06.rag.md",
        "07.resume-embedding-orcas.md",
        "08-질문생성.md"
    ]
    
    doc = Document()
    
    # 기본 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)

    # 제목 추가
    title = doc.add_heading('AI-워커 엔진 진행 보고서 (통합본)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for filename in files:
        file_path = os.path.join(base_path, filename)
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            continue
            
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # 파일 시작점 표시 (선택 사항이지만 구분용으로 추가)
        # doc.add_page_break() # 파일마다 페이지 나누고 싶다면 활성화
        doc.add_heading(f"--- {filename} ---", level=1)
        
        lines = content.split('\n')
        in_code_block = False
        code_text = ""
        
        for line in lines:
            # 코드 블록 처리
            if line.strip().startswith('```'):
                if in_code_block:
                    # 코드 블록 종료
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Pt(20)
                    in_code_block = False
                    code_text = ""
                else:
                    in_code_block = True
                continue
            
            if in_code_block:
                code_text += line + "\n"
                continue

            # 헤더 처리
            header_match = re.match(r'^(#+)\s+(.*)', line)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2)
                # Word는 레벨 1~9 지원
                doc.add_heading(text, level=min(level, 9))
                continue
            
            # 수평선 처리
            if re.match(r'^---|^===', line.strip()):
                doc.add_paragraph("_" * 50)
                continue
                
            # 일반 텍스트 및 리스트 (간단히)
            if line.strip():
                # 리스트인 경우 들여쓰기 효과
                if line.strip().startswith(('-', '*', '+')):
                    p = doc.add_paragraph(line.strip(), style='List Bullet')
                elif re.match(r'^\d+\.', line.strip()):
                    p = doc.add_paragraph(line.strip(), style='List Number')
                else:
                    doc.add_paragraph(line)
            else:
                # 빈 줄은 그냥 통과 (너무 많은 공백 방지)
                pass

    doc.save(output_filename)
    return output_filename

if __name__ == "__main__":
    # Get the root directory of the project (parent of "scripts")
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    base_dir = os.path.join(project_root, "파이널_진행보고서", "ai-워커")
    target_file = os.path.join(project_root, "AI_Worker_Comprehensive_Report_Combined.docx")
    try:
        markdown_to_docx(base_dir, target_file)
        print(f"Successfully created: {target_file}")
    except Exception as e:
        print(f"Error: {e}")
