
import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def single_markdown_to_docx(input_path, output_path):
    doc = Document()
    
    # 기본 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)

    if not os.path.exists(input_path):
        print(f"File not found: {input_path}")
        return
        
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    lines = content.split('\n')
    in_code_block = False
    code_text = ""
    
    for line in lines:
        # 코드 블록 처리
        if line.strip().startswith('```'):
            if in_code_block:
                # 코드 블록 종료
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Pt(20)
                in_code_block = False
                code_text = ""
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_text += line + "\n"
            continue

        # 헤더 처리
        header_match = re.match(r'^(#+)\s+(.*)', line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            # Word는 레벨 1~9 지원
            doc.add_heading(text, level=min(level, 9))
            continue
        
        # 수평선 처리
        if re.match(r'^---|^===', line.strip()):
            doc.add_paragraph("_" * 50)
            continue
            
        # 일반 텍스트 및 리스트 (간단히)
        if line.strip():
            # 리스트인 경우 들여쓰기 효과
            if line.strip().startswith(('-', '*', '+')):
                p = doc.add_paragraph(line.strip(), style='List Bullet')
            elif re.match(r'^\d+\.', line.strip()):
                p = doc.add_paragraph(line.strip(), style='List Number')
            else:
                doc.add_paragraph(line)
        else:
            # 빈 줄
            doc.add_paragraph()

    doc.save(output_path)
    print(f"Successfully converted: {os.path.basename(input_path)} -> {os.path.basename(output_path)}")

if __name__ == "__main__":
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    base_dir = os.path.join(project_root, "파이널_진행보고서", "ai-워커")
    
    files = [f for f in os.listdir(base_dir) if f.endswith('.md')]
    
    for filename in files:
        input_file = os.path.join(base_dir, filename)
        output_filename = filename.rsplit('.', 1)[0] + '.docx'
        output_file = os.path.join(base_dir, output_filename)
        try:
            single_markdown_to_docx(input_file, output_file)
        except Exception as e:
            print(f"Error converting {filename}: {e}")
