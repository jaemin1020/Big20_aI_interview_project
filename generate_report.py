
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('AI-Worker Engine Technical Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    doc.add_heading('1. 개요 (Project Overview)', level=1)
    doc.add_paragraph(
        "본 보고서는 AI 면접 시스템의 핵심 브레인 역할을 수행하는 'AI-Worker' 파트의 기술적 구현 사항과 알고리즘 설계를 상술합니다. "
        "기본적인 정적 시나리오를 넘어, 기업별 인재상(Ideal)을 DB에서 실시간으로 로드하여 지원자 맞춤형 질문을 생성하는 '지능형 동적 엔진' 구축을 핵심 성과로 합니다."
    )

    # Modules
    doc.add_heading('2. 모듈별 상세 분석 (System Architecture)', level=1)
    doc.add_paragraph("AI-Worker는 이력서 파싱, RAG 기반 검색, LLM 질문 생성, 평가 및 멀티모달 처리를 담당하는 독립적인 모듈들로 구성되어 있습니다.")
    
    # Table 1
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = '모듈명'
    hdr_cells[1].text = '주요 기능'
    hdr_cells[2].text = '적용 기술'
    hdr_cells[3].text = '데이터 I/O'
    
    data1 = [
        ['Resume Parser', '이력서 구조화', 'Layout Analysis, Zero-shot', 'PDF -> JSON'],
        ['RAG Engine', '문맥 정보 추출', 'Vector Search, pgvector', 'Query -> Context'],
        ['Question Gen', '맞춤형 질문 생성', 'LLM Prompting, Dynamic Var', 'Context -> Question'],
        ['Evaluator', '답변 분석 및 채점', 'Rubric Scoring, Sentiment', 'Transcript -> Score'],
        ['STT / TTS', '음성 변환', 'Whisper, TTS Synthesis', 'Audio <-> Text'],
        ['Vision AI', '비언어 분석', 'MediaPipe, Pose Est.', 'Video -> Data']
    ]
    for row in data1:
        row_cells = table1.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = val

    # RAG
    doc.add_heading('3. RAG 파이프라인 및 핵심 알고리즘', level=1)
    doc.add_paragraph("지원자의 이력서에서 가장 관련성 높은 정보를 실시간으로 추출하기 위해 고도화된 RAG(Retrieval-Augmented Generation) 프로세스를 적용했습니다.")
    
    # Table 2
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = '단계'
    hdr_cells[1].text = '처리 공정'
    hdr_cells[2].text = '핵심 메커니즘'
    
    data2 = [
        ['1. Chunking', '텍스트 분할', 'Recursive Splitting (Overlap 50)'],
        ['2. Embedding', '벡터 변환', 'Dense Vector (1024-dim)'],
        ['3. Indexing', '벡터 저장', 'pgvector HNSW Indexing'],
        ['4. Retrieval', '실시간 조회', 'Cosine Distance Scoring']
    ]
    for row in data2:
        row_cells = table2.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = val

    # Scenarios
    doc.add_heading('4. 기업 인재상 연동 및 동적 질문 전략', level=1)
    doc.add_paragraph(
        "DB의 'companies' 테이블에서 각 기업의 인재상(ideal) 데이터를 실시간으로 로드하여, "
        "AI가 이를 해석하고 질문에 반영하는 메커니즘을 구현했습니다."
    )
    
    # Table 3
    table3 = doc.add_table(rows=1, cols=3)
    table3.style = 'Table Grid'
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = '카테고리'
    hdr_cells[1].text = '생성 방식'
    hdr_cells[2].text = '인재상 연동 전략'
    
    data3 = [
        ['의사소통/협업', 'AI 생성 + 꼬리', '기술 공유 및 소통 능력 검증'],
        ['책임감/가치관', 'AI 생성 + 꼬리', '윤리적 책임 및 사회적 가치 딜레마'],
        ['성장의지/창의', 'AI 생성 + 꼬리', '새로운 관점 도입 및 융합 시도 사례']
    ]
    for row in data3:
        row_cells = table3.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = val

    # Infra
    doc.add_heading('5. 시스템 최적화 및 인프라', level=1)
    
    # Table 4
    table4 = doc.add_table(rows=1, cols=3)
    table4.style = 'Table Grid'
    hdr_cells = table4.rows[0].cells
    hdr_cells[0].text = '구분'
    hdr_cells[1].text = '적용 기술'
    hdr_cells[2].text = '설계 목적'
    
    data4 = [
        ['Task Queue', 'Celery / Redis', 'LLM 비동기 처리 (지연 방지)'],
        ['Vector DB', 'pgvector', '관계형/벡터 데이터 통합 관리'],
        ['Memory', 'PyTorch Cache', 'GPU 메모리 효율화 (GPU Leak 방지)'],
        ['Deployment', 'Docker Compose', '일관된 실행 환경 및 캐시 관리']
    ]
    for row in data4:
        row_cells = table4.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = val

    # Conclusion
    doc.add_heading('6. 결론', level=1)
    doc.add_paragraph(
        "이번 고도화를 통해 AI-Worker는 단순한 질문 생성기를 넘어 기업의 철학을 지원자의 역량과 매칭하는 '전략적 평가 엔진'으로 진화하였습니다. "
        "향후 실시간 데이터 분석 기능을 강화하여 더욱 정교한 면접 경험을 제공할 예정입니다."
    )

    file_path = r'C:\big20\Big20_aI_interview_project\AI_Worker_Technical_Report.docx'
    doc.save(file_path)
    return file_path

if __name__ == "__main__":
    path = create_report()
    print(f"Report created at: {path}")
