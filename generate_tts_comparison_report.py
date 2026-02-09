#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TTS 모델 비교 보고서 생성 스크립트
DOCX 파일로 Supertonic 2 vs Qwen3-TTS 비교 문서 생성
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_border(table):
    """표에 테두리 추가"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

# 문서 생성
doc = Document()

# 제목
title = doc.add_heading('TTS 모델 비교 분석 보고서', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 부제
subtitle = doc.add_paragraph('Supertonic 2 vs Qwen3-TTS')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()  # 빈 줄

# ============================================================
# 1. 최종 선택: Supertonic 2
# ============================================================
doc.add_heading('1. 최종 선택 모델', 1)

selection = doc.add_paragraph()
selection_run = selection.add_run('✅ Supertonic 2 (권장)')
selection_run.font.size = Pt(16)
selection_run.font.bold = True
selection_run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()

# 선택 이유
doc.add_heading('선택 이유', 2)
reasons = [
    '가볍고 빠른 성능 (1초 이내 생성)',
    '온프레미스(On-Premise) 방식으로 데이터 보안 강화',
    'ONNX Runtime 기반으로 저사양 환경에서도 안정적',
    '실시간 면접 시스템에 최적화된 응답 속도',
]

for reason in reasons:
    p = doc.add_paragraph(reason, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

# ============================================================
# 2. 모델 비교표
# ============================================================
doc.add_heading('2. 기능 비교', 1)

# 비교 표 생성 (5행 3열)
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
add_table_border(table)

# 헤더 행
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '기능'
hdr_cells[1].text = 'Supertonic 2 ⭐'
hdr_cells[2].text = 'Qwen3-TTS'

# 헤더 스타일
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 배경색 (연한 파란색)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'D9E2F3')
    cell._element.get_or_add_tcPr().append(shading_elm)

# 데이터 행
data_rows = [
    ['멘트 변경', '✅ 가능', '✅ 가능'],
    ['목소리 선택', '10개 (M1~M5, F1~F5)', '9개 (Vivian, Ethan 등)'],
    ['톤 조절', '❌ 불가능', '✅ 가능 (instruct로 감정/스타일 지시)'],
    ['속도', '⚡ 매우 빠름 (약 1초)', '🐢 느림 (약 37초)'],
]

for i, row_data in enumerate(data_rows, start=1):
    row_cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        cell = row_cells[j]
        cell.text = text
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        
        # 첫 번째 열 (기능명) 볼드
        if j == 0:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # 중앙 정렬
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Supertonic 열 (2번째) 배경색 강조
        if j == 1:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'E2EFDA')  # 연한 초록색
            cell._element.get_or_add_tcPr().append(shading_elm)

doc.add_paragraph()

# ============================================================
# 3. 상세 분석
# ============================================================
doc.add_heading('3. 상세 분석', 1)

# Supertonic 2 장점
doc.add_heading('Supertonic 2 주요 장점', 2)
advantages = [
    ('🚀 속도', '약 1초 이내로 음성 생성 완료 - 실시간 대화형 시스템에 이상적'),
    ('💾 모델 크기', '약 260MB로 경량화 - 서버 자원 효율적'),
    ('🔒 보안', '온프레미스 방식으로 외부 API 호출 없음 - 민감한 면접 데이터 보호'),
    ('🌐 한국어 지원', 'v1.6.0부터 공식 한국어 지원 - 자연스러운 발음'),
    ('⚙️ 안정성', 'ONNX Runtime 기반으로 크로스 플랫폼 지원'),
]

for title, desc in advantages:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    p.add_run(f': {desc}')
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

# Qwen3-TTS 장점
doc.add_heading('Qwen3-TTS 주요 장점', 2)
qwen_advantages = [
    ('🎭 감정 표현', 'instruct 파라미터로 다양한 톤 조절 가능 (친근함, 전문성, 진지함 등)'),
    ('🎤 목소리 품질', '프리미엄 화자 9개로 고급스러운 음질 제공'),
    ('🌍 다국어', '한국어 외 영어, 중국어, 일본어 등 10개 언어 지원'),
]

for title, desc in qwen_advantages:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    p.add_run(f': {desc}')
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

# ============================================================
# 4. 사용 사례별 추천
# ============================================================
doc.add_heading('4. 사용 사례별 추천', 1)

use_cases = [
    ('실시간 AI 면접 시스템', 'Supertonic 2', '빠른 응답 속도가 필수적'),
    ('저사양 서버 환경', 'Supertonic 2', '경량 모델로 자원 효율성 극대화'),
    ('데이터 보안이 중요한 경우', 'Supertonic 2', '온프레미스 방식으로 외부 유출 위험 없음'),
    ('다양한 감정 표현이 필요한 경우', 'Qwen3-TTS', '톤 조절 기능으로 맞춤형 응답 가능'),
    ('고급스러운 음질이 우선인 경우', 'Qwen3-TTS', '프리미엄 화자로 높은 품질 제공'),
]

# 표 생성
use_case_table = doc.add_table(rows=len(use_cases) + 1, cols=3)
use_case_table.style = 'Light List Accent 1'
add_table_border(use_case_table)

# 헤더
hdr = use_case_table.rows[0].cells
hdr[0].text = '사용 사례'
hdr[1].text = '추천 모델'
hdr[2].text = '이유'

for cell in hdr:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'FFF2CC')  # 연한 노란색
    cell._element.get_or_add_tcPr().append(shading_elm)

# 데이터
for i, (use_case, model, reason) in enumerate(use_cases, start=1):
    cells = use_case_table.rows[i].cells
    cells[0].text = use_case
    cells[1].text = model
    cells[2].text = reason
    
    # Supertonic 추천 행 강조
    if 'Supertonic' in model:
        cells[1].paragraphs[0].runs[0].font.bold = True
        cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()

# ============================================================
# 5. 결론
# ============================================================
doc.add_heading('5. 결론', 1)

conclusion = doc.add_paragraph()
conclusion.add_run('본 AI 면접 시스템에는 ').font.size = Pt(12)
run = conclusion.add_run('Supertonic 2')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 100, 0)
conclusion.add_run('를 최종 선택하였습니다. 실시간 응답이 필수적인 면접 환경에서 1초 이내의 빠른 속도와 온프레미스 방식의 보안성이 가장 큰 장점으로 작용했습니다.').font.size = Pt(12)

doc.add_paragraph()

# 추가 정보
info = doc.add_paragraph()
info.add_run('📝 작성일: 2026-02-08').font.size = Pt(10)
info.add_run('  |  ').font.size = Pt(10)
info.add_run('📌 프로젝트: AI 면접 시스템').font.size = Pt(10)
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 문서 저장
output_path = 'TTS_모델_비교_보고서.docx'
doc.save(output_path)
print(f"[OK] DOCX 파일 생성 완료: {output_path}")
