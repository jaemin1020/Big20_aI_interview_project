import sys
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    # Set default font to something that supports Korean well
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    font.size = Pt(10)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []

    for line in lines:
        line = line.strip()
        
        # Table handling
        if line.startswith('|'):
            if '---' in line: # Skip separator line
                continue
            in_table = True
            cells = [c.strip() for c in line.split('|') if c.strip()]
            table_data.append(cells)
            continue
        elif in_table:
            # End of table
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row_data in enumerate(table_data):
                    for j, cell_text in enumerate(row_data):
                        table.cell(i, j).text = cell_text
            in_table = False
            table_data = []

        if not line:
            continue

        # Horizontal Rule
        if line == '---' or line == '***':
            doc.add_page_break() # or just skip
            continue

        # Headings
        h_match = re.match(r'^(#+)\s+(.*)', line)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2)
            # Remove bold markup in headings
            text = text.replace('**', '')
            h = doc.add_heading(text, level=min(level, 9))
            continue

        # Lists
        l_match = re.match(r'^(\*|-)\s+(.*)', line)
        if l_match:
            text = l_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            continue

        n_match = re.match(r'^\d+\.\s+(.*)', line)
        if n_match:
            text = n_match.group(1)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            continue

        # Code blocks (simple)
        if line.startswith('```'):
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, line)

    # Save the document
    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

def add_formatted_text(paragraph, text):
    # Very simple bold parser
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    md_file = r"c:\big20\Big20_aI_interview_project\lyn\Final_Project_RAG_설계서.md"
    docx_file = r"c:\big20\Big20_aI_interview_project\lyn\Final_Project_RAG_설계서.docx"
    convert_md_to_docx(md_file, docx_file)
